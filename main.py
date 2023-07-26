import customtkinter
from tkinter import filedialog
from tkinter.filedialog import asksaveasfilename
import os
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import math

customtkinter.set_appearance_mode("dark")
customtkinter.set_default_color_theme("green")

# After selecting the workbook containing all of the data for your desired templates; capture the information in each meeting's worksheet
def captureTemplateData(file):
    '''
    Load in a selected workbook (be sure to use the mtg builder template I developed), loop through each of its sheets not named "TEMPLATE" and capturing the information on it.  
    Static information is the title, date, time, and facilitator, while attendees will require looping until we encounter a blank cell
    :Return meetings: List of Lists.  Each list contains static (header) data, and a list of attendees | each meeting - meeting[0] -> header meeting[1] -> attendees
    '''
    wb = load_workbook(file)

    meetings = [] # Empty list 

    for sheet in wb.worksheets:
        
        if sheet.title == 'TEMPLATE':
            continue

        # Capture our static data
        static_data = {
            'title' : sheet['b1'].value,
            'date' : sheet['b2'].value,
            'time' : sheet['b3'].value,
            'facil' : sheet['b4'].value
        }

        i = 1
        attendees = []
        print("Starting value in column C:", sheet[f'c{i}'].value )
        while sheet[f'c{i}'].value is not None:
            attendees.append(sheet[f'c{i}'].value)
            print("Added to our list:", sheet[f'c{i}'].value)
            i += 1

        meetings.append([static_data, attendees])

    return meetings

def buildTemplate(template, xl_file, note_bool = False, note_bool2 = False):

    def craftFileName(date, title, time):
        '''
        :Helper Function: builds the name of each file we are creating
        '''

        print("Original time", time)
        if 'pm' in time:
            ampm = 'pm'
        else:
            ampm = 'am'

        split_on_colon = time.split(":")
        lst = [split_on_colon[0]]
        
        if split_on_colon[1][0] == '0' and split_on_colon[1][1] == '0':
            lst.append(ampm)
        else:
            lst.append(split_on_colon[1][0] + split_on_colon[1][1])
            lst.append(ampm)

        mtg_time = "".join(lst)
        print("New built time", mtg_time)

        # Build our date
        splt_date = date.split('/')
        if len(splt_date[2]) == 4:
            yy = splt_date[2][2:]
        elif len(splt_date[2]) == 2:
            yy = splt_date[2]
        
        if len(splt_date[0]) == 1:
            mm = '0' + splt_date[0]
        else:
            mm = splt_date[0]

        if len(splt_date[1]) == 1:
            dd = '0' + splt_date[1]
        else:
            dd = splt_date[1]

        built_date = "".join([yy, mm, dd])
        
        # Build the filename: yymmdd_Mtg Notes_title_time.docx
        if note_bool is True:
            filename = built_date + "_Mtg Notes_" + title + "_" + mtg_time + ".docx"
        elif note_bool2 is True:
            filename = built_date + "_MTG NOTES_" + title + "_" + mtg_time + ".docx"
        else:
            filename = built_date + "_" + title + "_" + mtg_time + ".docx"
        return filename
    
    def updateAttendees(attendees, document):
        '''
        :Helper Function: Calculate the # of attendees/column, and then populate columns accordingly
        :Param attendees: List of people invited to the meeting
        :Param document: Template we open up to populate with each meeting's information - Should have a button for selecting which project this is for,
            so that the correct base_template is used to create each unique meeting template
        '''
        n = len(attendees) # number of attendees in our list

        attendees.sort(reverse=True) # Sort the attendee list so that we can pop and remove from the back-end in alphabetical order

        base = math.floor( n / 4 ) # number of attendees per column
        rem = n % 4 # add more attendees from left to right depending on remainder size

        # Determine how many attendees go in each column
        if rem == 0:
            colA, colB, colC, colD = base, base, base, base
        elif rem > 2:
            colD = base
            colA, colB, colC = base + 1, base + 1, base + 1
        elif rem > 1:
            colC, colD = base, base
            colA, colB = base + 1, base + 1
        elif rem > 0:
            colB, colC, colD = base, base, base
            colA = base + 1

        participant_table = document.tables[1] # access the participants table
        
        # "i" will represent the column we are on, we will go from left to right updating the odd columns with our attendees
        # we skip over the even columns because this is where we place an indicator for whether or not someone was present
        for i in range(8):
            if i == 1: 
                col = colA
            elif i == 3: 
                col = colB
            elif i == 5: 
                col = colC
            elif i == 7: 
                col = colD
            
            if i % 2 == 0:
                continue
            else:
                for k in range(col):
                    # Access our cell, add the appropriate name, and update font
                    cell = participant_table.cell(k, i)
                    cell.text = attendees.pop()
                    cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
                    cell.paragraphs[0].runs[0].font.size = Pt(11)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    def updateHeader(header_data, document):
        '''
        :Helper Function: Update the header of the loaded document to reflect static data captured in each worksheet
        :Param header_data: Informtation captured in each worksheet containing meeting information
        :Param document: Template for each meeting
        '''
                
        # Index the header table
        table = document.tables[0]

        # Update the header table with the date, facilitator, time, and title; and ensure we use the right font
        date_cell = table.cell(0, 3)
        date_cell.text = header_data['date']
        date_cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
        date_cell.paragraphs[0].runs[0].font.size = Pt(11)
        date_cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        date_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        date_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        if header_data['facil'] is None:
            print("No facilitator")
        else:
            facil_cell = table.cell(1, 1)
            facil_cell.text = header_data['facil']
            facil_cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
            facil_cell.paragraphs[0].runs[0].font.size = Pt(11)
            facil_cell.paragraphs[0].paragraph_format.line_spacing = 0
            facil_cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            facil_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            facil_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        time_cell = table.cell(1, 3)
        time_cell.text = header_data['time']
        time_cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
        time_cell.paragraphs[0].runs[0].font.size = Pt(11)
        time_cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        time_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        time_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        title_cell = table.cell(2, 1)
        title_cell.text = header_data['title']
        title_cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
        title_cell.paragraphs[0].runs[0].font.size = Pt(11)
        title_cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        title_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        title_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    meetings = captureTemplateData(xl_file) # grab each meeting's info from the loaded excel book

    for meeting in meetings: # Loop through each meeting object [header_data dict, attendees list] and build a template for it

        header_data = meeting[0]
        #print("Creating a template for:", header_data['title'], "at", header_data['time'], "on", header_data['date'])
        attendees = meeting[1] # Current meeting's attendees

        doc = Document(f'{template}') # Open up our project template as a word doc

        updateAttendees(attendees = attendees, document = doc) # Update the attendee table in our document

        updateHeader(header_data = header_data, document = doc) # Update the header in our document

        fname = craftFileName(header_data['date'], header_data['title'], header_data['time']) # Create the filename string

        # :Note: Make a function that creates a templates folder, navigates into it, then spits out each file
        doc.save(f'Templates/{fname}') # Save the updated template

class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        self.geometry("850x450")
        self.title("Template Builder")
        self.minsize(800, 300)

        # Create a 1x5 grid
        #self.grid_rowconfigure(0, weight = 1)
        self.grid_columnconfigure((0, 1, 2, 3, 4, 5, 6, 7), weight = 1)

        # Textbox: Display text to the user
        self.table_space = customtkinter.CTkTextbox(master=self, bg_color="white")
        self.table_space.grid(row = 0, column = 2, columnspan = 7, rowspan = 10, padx = 20, pady = (20, 0), sticky = "nsew")

        # Header label
        self.header_label = customtkinter.CTkLabel(master=self, text="Select Desired Project")
        self.header_label.grid(row = 0, column = 0, columnspan = 2, padx = 20, pady = (20, 0), sticky="w")

        # Project variable
        self.sport = customtkinter.StringVar()
        self.sport.set("FWB")

        # Buttons: Radiobuttons for each of our projects
        self.project_button_a = customtkinter.CTkRadioButton(master = self, text = "FWB", variable = self.sport, value = "FWB")
        self.project_button_b = customtkinter.CTkRadioButton(master = self, text = "CMP", variable = self.sport, value = "CMP")
        self.project_button_c = customtkinter.CTkRadioButton(master = self, text = "OE", variable = self.sport, value = "OE")
        self.project_button_d = customtkinter.CTkRadioButton(master = self, text = "SS", variable = self.sport, value = "SS")
        self.project_button_e = customtkinter.CTkRadioButton(master = self, text = "QH", variable = self.sport, value = "QH")
        self.project_button_f = customtkinter.CTkRadioButton(master = self, text = "SW", variable = self.sport, value = "SW")

        self.project_button_a.grid(row = 1, column = 0, padx = 20, pady = (10, 0), sticky="ew")
        self.project_button_b.grid(row = 1, column = 1, padx = 20, pady = (10, 0), sticky="ew")
        self.project_button_c.grid(row = 2, column = 0, padx = 20, pady = (10, 0), sticky="ew")
        self.project_button_d.grid(row = 2, column = 1, padx = 20, pady = (10, 0), sticky="ew")
        self.project_button_e.grid(row = 3, column = 0, padx = 20, pady = (10, 0), sticky="ew")
        self.project_button_f.grid(row = 3, column = 1, padx = 20, pady = (10, 0), sticky="ew")

        # Header: Download a fresh template
        self.data_label = customtkinter.CTkLabel(master=self, text="Download a fresh meeting data template")
        self.data_label.grid(row = 4, column = 0, padx = 20, pady = (20, 0), sticky="w")

        # Button: Download template
        self.master_file_button = customtkinter.CTkButton(master = self, text = "Download Template", command = self.pullTemplate)
        self.master_file_button.grid(row = 5, column = 0, rowspan=2, padx = 20, pady = 10, sticky="ew")

        # Header: Select data file
        self.data_label = customtkinter.CTkLabel(master=self, text="Select Data File")
        self.data_label.grid(row = 4, column = 1, padx = 20, pady = (20, 0), sticky="w")

        # Button: Select data (Excel) file (contains your meeting data)
        self.master_file_button = customtkinter.CTkButton(master = self, text = "Choose File", command = self.openFile)
        self.master_file_button.grid(row = 5, column = 1, rowspan=2, padx = 20, pady = 10, sticky="ew")

        # Variable: Mtg Notes in filename boolean 
        self.note_button_var = customtkinter.IntVar()
        self.note_button_var2 = customtkinter.IntVar()

        # Header: Mtg Notes label
        self.mtg_notes_checkbox_label = customtkinter.CTkLabel(master = self, text = 'Include "_Mtg Notes_" in the Filename')
        self.mtg_notes_checkbox_label.grid(row = 7, column = 0, padx = 20, pady = (20, 0), sticky="w")

        # Button:  Mtg Notes Checkbox (includes/excludes "_Mtg Notes_" from the filename)
        self.note_checkbox = customtkinter.CTkCheckBox(master = self, text="OFF", variable=self.note_button_var, command = self.toggle_button)
        self.note_checkbox.grid(row = 8, column = 0, padx = 20, pady = (10, 0), sticky="ew")

        # Header: Mtg Notes label
        self.mtg_notes_checkbox_label = customtkinter.CTkLabel(master = self, text = 'Include "_MTG NOTES_" in the Filename')
        self.mtg_notes_checkbox_label.grid(row = 7, column = 1, padx = 20, pady = (20, 0), sticky="w")

        # Button:  Mtg Notes Checkbox (includes/excludes "_Mtg Notes_" from the filename)
        self.note_checkbox2 = customtkinter.CTkCheckBox(master = self, text="OFF", variable=self.note_button_var2, command = self.toggle_button)
        self.note_checkbox2.grid(row = 8, column = 1, padx = 20, pady = (10, 0), sticky="ew")

        # Button:  Submit - Create your templates (must have selected a data file and project)
        self.submit_button = customtkinter.CTkButton(master = self, text = "Build Templates", command = self.Submit)
        self.submit_button.grid(row = 9, column = 1, padx = 20, pady = (30, 10), sticky="e")
    
    # Update our checkbox to display on/off depending on which is shown
    def toggle_button(self):
        if self.note_button_var.get() == 1:
            self.note_checkbox.configure(text="ON")
        else:
            self.note_checkbox.configure(text="OFF")

    # Update our checkbox to display on/off depending on which is shown
    def toggle_button(self):
        if self.note_button_var2.get() == 1:
            self.note_checkbox2.configure(text="ON")
        else:
            self.note_checkbox2.configure(text="OFF")
    
    # Method: Execute - Build meeting templates based on what's provided
    def Submit(self):
        # Execute the function that builds templates
        templates = {
        'FWB' : 'Base_Templates/fwb-template.docx',
        'OE' : 'Base_Templates/oe-template.docx', 
        'SS' : 'Base_Templates/ss-template.docx',
        'CMP' : 'Base_Templates/cmp-template.docx',
        'QH' : 'Base_Templates/qh-template.docx',
        'SW' : 'Base_Templates/sw-template.docx',
        }

        if self.note_button_var.get() == 1:
            note_toggle = True
        else:
            note_toggle = False

        if self.note_button_var2.get() == 1:
            note_toggle2 = True
        else:
            note_toggle2 = False

        sport = self.sport.get()
        project_template = templates.get(sport)

        buildTemplate(template = project_template, xl_file = self.master_file_name, note_bool = note_toggle, note_bool2 = note_toggle2)

    # Method: Download a fresh template
    def pullTemplate(self):
        filepath = asksaveasfilename(defaultextension='.xlsx', filetypes=[('Excel Files', '*.xlsx')])

        wb = load_workbook('Base_Templates/yymmdd_project_meeting builder.xlsx')
        wb.save(filepath)

    # Method: Ask the user which file they want to use to create meeting templates
    def openFile(self):
        
        file = filedialog.askopenfile(mode = 'r', filetypes = [("XLSX Files", "*.xlsx")])

        def captureFilename(string):
            '''
            Helper function that strips the location of the file off and just returns the filename
            '''
            i = -1
            ltrs = []
            while string[i] != '/': # add the letters of our file name (in reverse) to a list and stop when we reach the slash indicating that we've reached a folder
                ltrs.append(string[i])
                i -= 1
            
            s = []
            for i in range(len(ltrs)):
                s.append(ltrs.pop())

            return "".join(s)

        if file:
            filepath = os.path.abspath(file.name)
            #self.current_master_file = str(filepath)
            master_file_full_name = str(filepath)
            self.master_file_name = captureFilename(master_file_full_name)
            self.master_file_text = customtkinter.CTkLabel(master = self, text = "Selected File: " + self.master_file_name)
            self.master_file_text.grid(row=10, column = 0, columnspan=8, padx = 10, pady = (20, 0))
            self.open_file_bool = True
            #self.master_file_text.text = str(filepath)

if __name__ == "__main__":
    app = App()
    app.mainloop()