import customtkinter
import tkinter
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import math
from openpyxl import load_workbook
import os

# After selecting the workbook containing all of the data for your desired templates; capture the information in each meeting's worksheet
def captureTemplateData(file):
    '''
    Load in a selected workbook (be sure to use the mtg builder template I developed), loop through each of its sheets not named "TEMPLATE" and capturing the information on it.  
    Static information is the title, date, time, and facilitator, while attendees will require looping until we encounter a blank cell
    :Return meetings: List of Lists.  Each list contains static (header) data, and a list of attendees | each meeting - meeting[0] -> header meeting[1] -> attendees
    '''
    wb = load_workbook(file)

    meetings = [] # Empty list 

    for sheet in wb.worksheets:
        
        if sheet.title == 'TEMPLATE':
            continue

        print("Currently accessing worksheet", sheet.title)
        # Capture our static data
        static_data = {
            'title' : sheet['b1'].value,
            'date' : sheet['b2'].value,
            'time' : sheet['b3'].value,
            'facil' : sheet['b4'].value
        }
        print("Static data captured:", static_data)

        i = 1
        attendees = []
        print("Starting value in column C:", sheet[f'c{i}'].value )
        while sheet[f'c{i}'].value is not None:
            attendees.append(sheet[f'c{i}'].value)
            print("Added to our list:", sheet[f'c{i}'].value)
            i += 1
        print("Finished capturing information for", sheet.title, "\n=======================================")

        meetings.append([static_data, attendees])

    return meetings

def buildTemplate(template, xl_file):

    def craftFileName(date, title, time):
        '''
        :Helper Function: builds the name of each file we are creating
        '''
        # Remove ":" from our time field
        splt_time = time.split(":")
        mtg_time = "".join(splt_time)

        # Build our date
        splt_date = date.split('/')
        if len(splt_date[2]) == 4:
            yy = splt_date[2][2:]
        elif len(splt_date[2]) == 2:
            yy = splt_date[2]
        
        if len(splt_date[0]) == 1:
            mm = '0' + splt_date[0]
        else:
            mm = splt_date[0]

        if len(splt_date[1]) == 1:
            dd = '0' + splt_date[1]
        else:
            dd = splt_date[1]

        built_date = "".join([yy, mm, dd])

        # remove space from our meeting time, if it was enterred into the template
        try:
            mtg_time = mtg_time.split(' ')
            mtg_time = "".join(mtg_time)
        except:
            pass
        
        # Build the filename: yymmdd_Mtg Notes_title_time.docx
        filename = built_date + "_" + title + "_" + mtg_time + ".docx"
        return filename
    
    def updateAttendees(attendees, document):
        '''
        :Helper Function: Calculate the # of attendees/column, and then populate columns accordingly
        :Param attendees: List of people invited to the meeting
        :Param document: Template we open up to populate with each meeting's information - Should have a button for selecting which project this is for,
            so that the correct base_template is used to create each unique meeting template
        '''
        n = len(attendees) # number of attendees in our list

        attendees.sort(reverse=True) # Sort the attendee list so that we can pop and remove from the back-end in alphabetical order

        base = math.floor( n / 4 ) # number of attendees per column
        rem = n % 4 # add more attendees from left to right depending on remainder size

        # Determine how many attendees go in each column
        if rem == 0:
            colA, colB, colC, colD = base, base, base, base
        elif rem > 2:
            colD = base
            colA, colB, colC = base + 1, base + 1, base + 1
        elif rem > 1:
            colC, colD = base, base
            colA, colB = base + 1, base + 1
        elif rem > 0:
            colB, colC, colD = base, base, base
            colA = base + 1

        participant_table = document.tables[1] # access the participants table
        
        # "i" will represent the column we are on, we will go from left to right updating the odd columns with our attendees
        # we skip over the even columns because this is where we place an indicator for whether or not someone was present
        for i in range(8):
            if i == 1: 
                col = colA
            elif i == 3: 
                col = colB
            elif i == 5: 
                col = colC
            elif i == 7: 
                col = colD
            
            if i % 2 == 0:
                continue
            else:
                for k in range(col):
                    # Access our cell, add the appropriate name, and update font
                    cell = participant_table.cell(k, i)
                    cell.text = attendees.pop()
                    cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
                    cell.paragraphs[0].runs[0].font.size = Pt(11)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    def updateHeader(header_data, document):
        '''
        :Helper Function: Update the header of the loaded document to reflect static data captured in each worksheet
        :Param header_data: Informtation captured in each worksheet containing meeting information
        :Param document: Template for each meeting
        '''
                
        # Index the header table
        table = document.tables[0]

        # Update the header table with the date, facilitator, time, and title; and ensure we use the right font
        date_cell = table.cell(0, 3)
        date_cell.text = header_data['date']
        date_cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
        date_cell.paragraphs[0].runs[0].font.size = Pt(11)
        date_cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        date_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        date_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        facil_cell = table.cell(1, 1)
        facil_cell.text = header_data['facil']
        facil_cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
        facil_cell.paragraphs[0].runs[0].font.size = Pt(11)
        facil_cell.paragraphs[0].paragraph_format.line_spacing = 0
        facil_cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        facil_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        facil_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        time_cell = table.cell(1, 3)
        time_cell.text = header_data['time']
        time_cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
        time_cell.paragraphs[0].runs[0].font.size = Pt(11)
        time_cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        time_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        time_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        title_cell = table.cell(2, 1)
        title_cell.text = header_data['title']
        title_cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
        title_cell.paragraphs[0].runs[0].font.size = Pt(11)
        title_cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        title_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        title_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    meetings = captureTemplateData(xl_file) # grab each meeting's info from the loaded excel book

    for meeting in meetings: # Loop through each meeting object [header_data dict, attendees list] and build a template for it

        header_data = meeting[0]
        #print("Header Data", header_data)
        print("Creating a template for:", header_data['title'], "at", header_data['time'], "on", header_data['date'])
        attendees = meeting[1] # Current meeting's attendees
        #print("Including the following attendees:", attendees)

        doc = Document(f'{template}') # Open up our project template as a word doc

        updateAttendees(attendees = attendees, document = doc) # Update the attendee table in our document

        updateHeader(header_data = header_data, document = doc) # Update the header in our document

        fname = craftFileName(header_data['date'], header_data['title'], header_data['time']) # Create the filename string


        # :Note: Make a function that creates a templates folder, navigates into it, then spits out each file
        doc.save(f'Templates/{fname}') # Save the updated template

        print(f"{fname} saved!")

# project_template = 'sw-template.docx' # Which base_template do you want to use? - Create mappings for dif projects and dif base_template files

templates = {
    '1' : 'Base_Templates/fwb-template.docx',
    '2' : 'Base_Templates/oe-template.docx', 
    '3' : 'Base_Templates/sw-template.docx',
    '4' : 'Base_Templates/cmp-template.docx',
    '5' : 'Base_Templates/qh-template.docx',
}

user_inp = input("Enter 1 for FWB, 2 for OE, 3 for SW, 4 for CMP, 5 for QH: ")

project_template = templates.get(user_inp)

xl_file = 'mtg_builder_4-28-qh.xlsx' # Which file contains our template information?

buildTemplate(template = project_template, xl_file = xl_file)

# I should have a log so that with each file we create, we also create a text file saying which meetings blah blah blah